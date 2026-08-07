"""
output_formatters.py
Plain-text output writers for the Khmer OCR pipeline.

Segment schema
--------------
  {"type": "text", "text": str,        "bbox": (x1,y1,x2,y2)}
  {"type": "logo", "crop": PIL.Image,  "bbox": (x1,y1,x2,y2)}

Supported formats (auto-detected from file extension):
  .txt   — text segments only, one line per OCR line
  .md    — same as .txt
  .json  — text segments only, structured metadata (image info + text + bbox)
  .docx  — text as paragraphs, logos as inline images
"""

import io
import json
import os

SUPPORTED_FORMATS = {".txt", ".md", ".json", ".docx"}


def save_output(
    segments: list,
    output_path: str,
    image_path: str | None = None,
    image_size: tuple | None = None,
) -> None:
    """
    Write OCR results to file.

    Args:
        segments: list of segment dicts (type "text" or "logo")
        output_path: destination file; extension selects the format
        image_path: source image path (used by JSON output)
        image_size: (width, height) of the source image (used by JSON + DOCX)
    """
    ext = os.path.splitext(output_path)[1].lower()
    dispatch = {
        ".txt":  _save_txt,
        ".md":   _save_md,
        ".json": _save_json,
        ".docx": _save_docx,
    }
    fn = dispatch.get(ext)
    if fn is None:
        raise ValueError(
            f"Unsupported format '{ext}'. Supported: {sorted(SUPPORTED_FORMATS)}"
        )
    fn(segments, output_path, image_path=image_path, image_size=image_size)
    print(f"  Saved {ext} output -> {output_path}")


# ──────────────────────────────────────────────────────────────────────
# Plain-text (text segments only)
# ──────────────────────────────────────────────────────────────────────

def _save_txt(segments, output_path, **_):
    text = "\n".join(s["text"] for s in segments if s.get("type") == "text")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)


def _save_md(segments, output_path, **_):
    text = "\n".join(s["text"] for s in segments if s.get("type") == "text")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)


# ──────────────────────────────────────────────────────────────────────
# JSON — text segments only (logos are images, not representable in JSON)
# ──────────────────────────────────────────────────────────────────────

def _save_json(segments, output_path, image_path=None, image_size=None, **_):
    payload = {
        "image_path": image_path,
        "image_size": list(image_size) if image_size else None,
        "lines": [
            {"text": s["text"], "bbox": list(s["bbox"])}
            for s in segments if s.get("type") == "text"
        ],
    }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)


# ──────────────────────────────────────────────────────────────────────
# DOCX — text as paragraphs, logos as inline images
# ──────────────────────────────────────────────────────────────────────

def _save_docx(segments, output_path, image_size=None, **_):
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as e:
        raise ImportError(
            "Writing .docx output requires the 'docx' extra. Install it with: "
            "pip install \"netra-ocr[docx]\""
        ) from e

    doc = Document()
    page_width_in = 6.0                          # usable page width in inches
    orig_w = image_size[0] if image_size else None

    for seg in segments:
        if seg.get("type") == "logo":
            buf = io.BytesIO()
            seg["crop"].save(buf, format="PNG")
            buf.seek(0)
            crop_w = seg["crop"].size[0]
            if orig_w:
                width_in = (crop_w / orig_w) * page_width_in
            else:
                width_in = min(page_width_in, crop_w / 150)  # fallback: 150 DPI
            doc.add_picture(buf, width=Inches(min(page_width_in, width_in)))
        else:
            doc.add_paragraph(seg["text"])

    doc.save(output_path)
